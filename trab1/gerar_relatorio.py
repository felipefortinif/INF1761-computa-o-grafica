from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Margens
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Estilo base
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    p.runs[0].bold = True
    return p

def para(doc, text, justify=True):
    p = doc.add_paragraph(text)
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def mono(run):
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

# ── Título ──────────────────────────────────────────────────────────────────
t = doc.add_paragraph('Relógio Analógico com WebGPU')
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.runs[0]; r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'

s = doc.add_paragraph('Trabalho Prático – INF1761 Computação Gráfica')
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.runs[0].font.size = Pt(11); s.runs[0].font.name = 'Times New Roman'

i = doc.add_paragraph('Felipe Fortini  |  PUC-Rio  |  2026')
i.alignment = WD_ALIGN_PARAGRAPH.CENTER
i.runs[0].font.size = Pt(11); i.runs[0].font.name = 'Times New Roman'
i.runs[0].font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# ── 1. Introdução ────────────────────────────────────────────────────────────
heading(doc, '1. Introdução')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run(
    'Este trabalho implementa um relógio analógico interativo utilizando a API WebGPU, '
    'executado diretamente no navegador. O relógio exibe a hora corrente do sistema, '
    'com ponteiros de hora, minuto e segundo atualizados dinamicamente a cada quadro renderizado.'
)
r.font.name = 'Times New Roman'; r.font.size = Pt(12)

# ── 2. Implementação ─────────────────────────────────────────────────────────
heading(doc, '2. Implementação')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run(
    'O programa é composto por um único arquivo '
); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r2 = p.add_run('index.html'); mono(r2)
r3 = p.add_run(
    ' com código JavaScript e shaders escritos em WGSL (WebGPU Shading Language). '
    'A estrutura geral segue o pipeline padrão de WebGPU: criação do adaptador e dispositivo, '
    'configuração do contexto do canvas, definição do pipeline de renderização e loop de '
    'animação via '
); r3.font.name = 'Times New Roman'; r3.font.size = Pt(12)
r4 = p.add_run('requestAnimationFrame'); mono(r4)
r5 = p.add_run('.'); r5.font.name = 'Times New Roman'; r5.font.size = Pt(12)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p2.add_run(
    'A geometria é representada por triângulos (topologia '
); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r2 = p2.add_run('triangle-list'); mono(r2)
r3 = p2.add_run(
    '), única primitiva necessária para compor todas as formas do relógio. '
    'Os vértices carregam posição ('
); r3.font.name = 'Times New Roman'; r3.font.size = Pt(12)
r4 = p2.add_run('vec2f'); mono(r4)
r5 = p2.add_run(') e cor ('); r5.font.name = 'Times New Roman'; r5.font.size = Pt(12)
r6 = p2.add_run('vec4f'); mono(r6)
r7 = p2.add_run(
    '), totalizando 6 floats por vértice. A conversão de coordenadas em pixels para '
    'clip-space é feita no vertex shader usando a resolução do canvas passada via uniform buffer.'
); r7.font.name = 'Times New Roman'; r7.font.size = Pt(12)

# ── 3. Geometria ─────────────────────────────────────────────────────────────
heading(doc, '3. Geometria das Formas')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run(
    'A tabela abaixo descreve os elementos visuais do relógio e as formas geométricas utilizadas.'
); r.font.name = 'Times New Roman'; r.font.size = Pt(12)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
hdr[0].text = 'Elemento'
hdr[1].text = 'Forma geométrica'
hdr[2].text = 'Descrição'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell.paragraphs[0].runs[0].font.size = Pt(11)

rows_data = [
    ('Face do relógio',          'Disco circular',        'Fan de 64 triângulos a partir do centro'),
    ('Borda',                    'Anel (annulus)',         '64 pares de triângulos entre raio interno e externo'),
    ('Marcações de hora',        'Retângulos',            '12 retângulos finos rotacionados, 2 triângulos cada'),
    ('Ponteiros de hora e min.', 'Quadrilátero (trapézio)','Polígono de 4 vértices (2 triângulos); base larga e ponta estreita'),
    ('Ponteiro de segundo',      'Quadrilátero fino',     'Mesmo esquema, largura reduzida e cor vermelha'),
    ('Pino central',             'Disco pequeno',         'Fan de 16 triângulos sobre o centro'),
]
for rd in rows_data:
    row = table.add_row().cells
    for j, val in enumerate(rd):
        row[j].text = val
        row[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row[j].paragraphs[0].runs[0].font.size = Pt(11)

doc.add_paragraph()

# ── 4. Animação ──────────────────────────────────────────────────────────────
heading(doc, '4. Animação')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run(
    'A cada quadro, a hora corrente é lida via '
); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r2 = p.add_run('new Date()'); mono(r2)
r3 = p.add_run(
    '. Os ângulos dos ponteiros são calculados com precisão de milissegundos, '
    'garantindo movimento suave:'
); r3.font.name = 'Times New Roman'; r3.font.size = Pt(12)

bullets = [
    'Segundo:  (s + ms/1000) / 60 × 2π',
    'Minuto:   (m + s/60) / 60 × 2π',
    'Hora:     (h + m/60) / 12 × 2π',
]
for b in bullets:
    bp = doc.add_paragraph(style='List Bullet')
    r = bp.add_run(b); r.font.name = 'Times New Roman'; r.font.size = Pt(12)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p2.add_run(
    'A geometria dos ponteiros é reconstruída a cada frame e enviada à GPU via '
); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r2 = p2.add_run('writeBuffer'); mono(r2)
r3 = p2.add_run(
    ' em um buffer dinâmico separado do buffer estático (face, borda e marcações), '
    'que é gravado uma única vez na inicialização.'
); r3.font.name = 'Times New Roman'; r3.font.size = Pt(12)

# ── 5. Resultado ─────────────────────────────────────────────────────────────
heading(doc, '5. Resultado')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run(
    'O programa renderiza corretamente o relógio analógico com os três ponteiros '
    'sincronizados ao horário local do sistema. O vídeo demonstrando o funcionamento '
    'está disponível em:'
); r.font.name = 'Times New Roman'; r.font.size = Pt(12)

lp = doc.add_paragraph()
lp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
rb = lp.add_run('Link do vídeo: '); rb.bold = True
rb.font.name = 'Times New Roman'; rb.font.size = Pt(12)
rl = lp.add_run('INSERIR_LINK_AQUI')
rl.font.name = 'Times New Roman'; rl.font.size = Pt(12)

# ── 6. Tecnologias ───────────────────────────────────────────────────────────
heading(doc, '6. Tecnologias Utilizadas')
techs = [
    'WebGPU (API nativa do navegador)',
    'WGSL – WebGPU Shading Language (vertex e fragment shaders)',
    'JavaScript (ES2022, módulos nativos)',
    'HTML5 Canvas',
]
for t in techs:
    bp = doc.add_paragraph(style='List Bullet')
    r = bp.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('relatorio.docx')
print('relatorio.docx gerado com sucesso.')
